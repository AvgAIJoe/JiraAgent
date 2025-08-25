import streamlit as st
import json
import pandas as pd
from io import StringIO
import docx
from openai import OpenAI

st.set_page_config(
    page_title="Jira-Agent",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Clean professional styling
st.markdown("""
<style>
    /* Main header styling */
    .main-header {
        background: linear-gradient(90deg, #0066cc 0%, #004499 100%);
        color: white;
        padding: 2rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        text-align: center;
        box-shadow: 0 2px 10px rgba(0, 102, 204, 0.2);
    }
    
    .main-header h1 {
        margin: 0;
        font-size: 2.5rem;
        font-weight: 600;
        letter-spacing: -0.5px;
    }
    
    .main-header p {
        margin: 0.5rem 0 0 0;
        opacity: 0.95;
        font-size: 1.1rem;
    }
    
    /* Hide default streamlit elements */
    #MainMenu {visibility: hidden;}
    .stDeployButton {display: none;}
    footer {visibility: hidden;}
    
    /* Better spacing */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    
    /* Card styling */
    .stExpander {
        border: 1px solid #e1e5e9;
        border-radius: 8px;
        margin-bottom: 0.5rem;
    }
    
    /* Epic styling */
    .epic-title {
        background: #e8f4fd;
        border-left: 4px solid #0066cc;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        font-weight: 600;
        color: #004499;
    }
</style>
""", unsafe_allow_html=True)

def trigger_processing():
    st.session_state.should_process = True

def main():
    # Initialize session state
    if 'should_process' not in st.session_state:
        st.session_state.should_process = False
    if 'last_content' not in st.session_state:
        st.session_state.last_content = ""
    if 'last_file' not in st.session_state:
        st.session_state.last_file = None
    
    # Professional header
    st.markdown("""
    <div class="main-header">
        <h1>Jira-Agent</h1>
        <p>Transform requirements into structured JIRA tickets with AI</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        st.divider()
        
        api_key = st.text_input(
            "OpenAI API Key", 
            type="password", 
            help="Enter your OpenAI API key",
            placeholder="sk-..."
        )
        
        if api_key:
            st.success("✅ API Key configured")
        else:
            st.info("🔑 Enter API key to get started")
            
        st.divider()
        st.subheader("Model Settings")
        
        model_choice = st.selectbox(
            "AI Model",
            ["gpt-4o-mini", "gpt-4", "gpt-3.5-turbo"],
            help="Choose the AI model for processing"
        )
        
        temperature = st.slider(
            "Creativity Level",
            min_value=0.0,
            max_value=1.0,
            value=0.7,
            step=0.1,
            help="Higher values make output more creative"
        )
    
    # Input section
    st.header("📝 Input Requirements")
    
    tab1, tab2 = st.tabs(["📄 Upload Document", "📝 Paste Text"])
    
    # Initialize variables outside of tabs
    uploaded_file = None
    text_input = ""
    
    with tab1:
        uploaded_file = st.file_uploader(
            "Upload Word Document",
            type=['docx'],
            help="Upload a .docx file containing your requirements"
        )
        
        if uploaded_file:
            st.success(f"✅ File uploaded: {uploaded_file.name}")
            file_size = round(uploaded_file.size / 1024, 1)
            st.caption(f"File size: {file_size} KB")
    
    with tab2:
        text_input = st.text_area(
            "Paste Requirements",
            height=300,
            placeholder="Paste your requirements, features, or project description here...\n\nExample:\n• User Management System\n• Authentication and Authorization  \n• Dashboard with Analytics\n• Mobile Responsive Design"
        )
        
        if text_input:
            word_count = len(text_input.split())
            st.caption(f"Word count: {word_count}")
    
    # Store current content in session state for processing
    current_content = None
    if uploaded_file is not None:
        current_content = uploaded_file.name  # Use filename as identifier
        st.session_state.current_file = uploaded_file
    elif text_input and text_input.strip():
        current_content = text_input.strip()
        st.session_state.current_text = text_input.strip()
    
    # Processing button
    st.divider()
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.button(
            "🚀 Generate JIRA Tickets", 
            type="primary", 
            use_container_width=True,
            on_click=trigger_processing
        )
    
    # Process when session state indicates processing should happen
    if st.session_state.should_process:
        # Reset the flag
        st.session_state.should_process = False
        
        if not api_key:
            st.error("🔑 Please provide your OpenAI API key in the sidebar")
            st.stop()
        
        # Get content from session state
        content = None
        if hasattr(st.session_state, 'current_file') and st.session_state.current_file is not None:
            content = extract_text_from_docx(st.session_state.current_file)
        elif hasattr(st.session_state, 'current_text') and st.session_state.current_text:
            content = st.session_state.current_text
        
        if not content:
            st.error("📝 Please provide input either by uploading a document or pasting text")
            st.stop()
        
        # Process the content
        with st.spinner("🤖 AI is analyzing your requirements..."):
            try:
                structured_data = process_with_openai(content, api_key, model_choice, temperature)
                display_results(structured_data)
                
            except Exception as e:
                st.error(f"❌ Error processing requirements: {str(e)}")

def extract_text_from_docx(file):
    """Extract text content from uploaded DOCX file"""
    try:
        doc = docx.Document(file)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
    
    except Exception as e:
        st.error(f"Error reading document: {str(e)}")
        return None

def process_with_openai(content, api_key, model="gpt-4o-mini", temperature=0.7):
    """Process content using OpenAI to generate structured JIRA tickets"""
    client = OpenAI(api_key=api_key)
    
    prompt = f"""
    Analyze the following requirements document and extract structured information to create JIRA tickets.
    
    Please break down the content into:
    1. Epics (high-level features or themes)
    2. Stories (user-facing functionality)
    3. Sub-tasks (technical implementation tasks)
    
    For each item, provide:
    - Title (clear and concise)
    - Description (detailed explanation)
    - Acceptance Criteria (specific conditions for completion)
    - Story Point Estimate (1, 2, 3, 5, 8, 13)
    
    Return the result as a JSON object with this structure:
    {{
        "epics": [
            {{
                "title": "Epic title",
                "description": "Epic description",
                "estimate": 13,
                "stories": [
                    {{
                        "title": "Story title",
                        "description": "Story description",
                        "acceptance_criteria": "Acceptance criteria",
                        "estimate": 5,
                        "subtasks": [
                            {{
                                "title": "Subtask title",
                                "description": "Subtask description",
                                "acceptance_criteria": "Acceptance criteria",
                                "estimate": 2
                            }}
                        ]
                    }}
                ]
            }}
        ]
    }}
    
    Content to analyze:
    {content}
    """
    
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are an expert business analyst and project manager who specializes in breaking down requirements into structured JIRA tickets."},
            {"role": "user", "content": prompt}
        ],
        temperature=temperature
    )
    
    # Parse JSON response
    try:
        return json.loads(response.choices[0].message.content)
    except json.JSONDecodeError:
        content = response.choices[0].message.content
        start = content.find('{')
        end = content.rfind('}') + 1
        if start != -1 and end > start:
            return json.loads(content[start:end])
        else:
            raise ValueError("Could not parse JSON from OpenAI response")

def display_results(data):
    """Display the structured results"""
    st.divider()
    st.header("📋 Generated JIRA Structure")
    
    # JSON Preview
    with st.expander("🔍 Raw JSON Data", expanded=False):
        st.json(data)
    
    # Display structure
    total_story_points = 0
    
    for epic_idx, epic in enumerate(data.get('epics', [])):
        # Epic section
        st.markdown(f"""<div class="epic-title">🎯 Epic {epic_idx + 1}: {epic.get('title', 'Untitled Epic')}</div>""", unsafe_allow_html=True)
        
        col1, col2 = st.columns([3, 1])
        with col1:
            st.write(f"**Description:** {epic.get('description', 'No description')}")
        with col2:
            epic_estimate = epic.get('estimate', 0)
            st.metric("Story Points", epic_estimate)
            total_story_points += epic_estimate
        
        # Stories
        for story_idx, story in enumerate(epic.get('stories', [])):
            with st.expander(f"📖 Story {story_idx + 1}: {story.get('title', 'Untitled Story')}"):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(f"**Description:** {story.get('description', 'No description')}")
                    st.write(f"**Acceptance Criteria:** {story.get('acceptance_criteria', 'No AC defined')}")
                with col2:
                    story_estimate = story.get('estimate', 0)
                    st.metric("Story Points", story_estimate)
                    total_story_points += story_estimate
                
                # Subtasks
                if story.get('subtasks', []):
                    st.write("**Subtasks:**")
                    for subtask_idx, subtask in enumerate(story.get('subtasks', [])):
                        st.markdown(f"**⚙️ Subtask {subtask_idx + 1}: {subtask.get('title', 'Untitled Subtask')}**")
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            st.write(f"*Description:* {subtask.get('description', 'No description')}")
                            st.write(f"*Acceptance Criteria:* {subtask.get('acceptance_criteria', 'No AC defined')}")
                        with col2:
                            subtask_estimate = subtask.get('estimate', 0)
                            st.metric("Story Points", subtask_estimate)
                            total_story_points += subtask_estimate
                        st.write("")  # Add spacing between subtasks
        
        st.write("")
    
    # Summary
    st.divider()
    st.header("📊 Project Summary")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Epics", len(data.get('epics', [])))
    with col2:
        total_stories = sum(len(epic.get('stories', [])) for epic in data.get('epics', []))
        st.metric("Stories", total_stories)
    with col3:
        total_subtasks = sum(len(subtask) for epic in data.get('epics', []) for subtask in [story.get('subtasks', []) for story in epic.get('stories', [])])
        st.metric("Subtasks", total_subtasks)
    with col4:
        st.metric("Story Points", total_story_points)
    
    # Export options
    st.divider()
    st.header("📤 Export Options")
    
    col1, col2, col3 = st.columns(3)
    
    json_str = json.dumps(data, indent=2)
    csv_data = convert_to_csv(data)
    text_format = convert_to_text(data)
    
    with col1:
        st.download_button(
            label="Download JSON",
            data=json_str,
            file_name=f"jira_tickets_{pd.Timestamp.now().strftime('%Y%m%d_%H%M')}.json",
            mime="application/json",
            use_container_width=True
        )
    
    with col2:
        st.download_button(
            label="Download CSV",
            data=csv_data,
            file_name=f"jira_tickets_{pd.Timestamp.now().strftime('%Y%m%d_%H%M')}.csv",
            mime="text/csv",
            use_container_width=True
        )
    
    with col3:
        st.download_button(
            label="Download Text",
            data=text_format,
            file_name=f"jira_tickets_{pd.Timestamp.now().strftime('%Y%m%d_%H%M')}.txt",
            mime="text/plain",
            use_container_width=True
        )

def convert_to_csv(data):
    """Convert structured data to CSV format"""
    rows = []
    
    for epic in data.get('epics', []):
        rows.append({
            'Type': 'Epic',
            'Title': epic.get('title', ''),
            'Description': epic.get('description', ''),
            'Acceptance Criteria': '',
            'Story Points': epic.get('estimate', 0),
            'Parent': ''
        })
        
        for story in epic.get('stories', []):
            rows.append({
                'Type': 'Story',
                'Title': story.get('title', ''),
                'Description': story.get('description', ''),
                'Acceptance Criteria': story.get('acceptance_criteria', ''),
                'Story Points': story.get('estimate', 0),
                'Parent': epic.get('title', '')
            })
            
            for subtask in story.get('subtasks', []):
                rows.append({
                    'Type': 'Sub-task',
                    'Title': subtask.get('title', ''),
                    'Description': subtask.get('description', ''),
                    'Acceptance Criteria': subtask.get('acceptance_criteria', ''),
                    'Story Points': subtask.get('estimate', 0),
                    'Parent': story.get('title', '')
                })
    
    df = pd.DataFrame(rows)
    return df.to_csv(index=False)

def convert_to_text(data):
    """Convert structured data to readable text format"""
    text_parts = []
    
    for epic_idx, epic in enumerate(data.get('epics', [])):
        text_parts.append(f"EPIC {epic_idx + 1}: {epic.get('title', 'Untitled Epic')}")
        text_parts.append(f"Description: {epic.get('description', 'No description')}")
        text_parts.append(f"Story Points: {epic.get('estimate', 0)}")
        text_parts.append("")
        
        for story_idx, story in enumerate(epic.get('stories', [])):
            text_parts.append(f"  STORY {story_idx + 1}: {story.get('title', 'Untitled Story')}")
            text_parts.append(f"  Description: {story.get('description', 'No description')}")
            text_parts.append(f"  Acceptance Criteria: {story.get('acceptance_criteria', 'No AC defined')}")
            text_parts.append(f"  Story Points: {story.get('estimate', 0)}")
            text_parts.append("")
            
            for subtask_idx, subtask in enumerate(story.get('subtasks', [])):
                text_parts.append(f"    SUBTASK {subtask_idx + 1}: {subtask.get('title', 'Untitled Subtask')}")
                text_parts.append(f"    Description: {subtask.get('description', 'No description')}")
                text_parts.append(f"    Acceptance Criteria: {subtask.get('acceptance_criteria', 'No AC defined')}")
                text_parts.append(f"    Story Points: {subtask.get('estimate', 0)}")
                text_parts.append("")
        
        text_parts.append("-" * 50)
        text_parts.append("")
    
    return "\n".join(text_parts)

if __name__ == "__main__":
    main()